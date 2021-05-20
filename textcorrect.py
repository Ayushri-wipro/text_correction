import re
import os
import docx
import requests
from transformers import AutoTokenizer, T5ForConditionalGeneration
import flask
from werkzeug.utils import secure_filename
import codecs

app = flask.Flask(__name__)
app.config["DEBUG"] = True

#UPLOAD_FOLDER = '/home/ayushri/Desktop'

@app.route('/', methods=['GET'])
def check():
    return "<h1>In Textcorrector container</h1>"


@app.route('/api/v1/getfile/', methods=['POST'])
def getfile():
    content= flask.request.data
    #print(content)
    str= codecs.decode(content)
    #filepath = flask.request.files['file'].read()
    #filename= filepath.filename
    #filestream= filepath.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
    #doc = docx.Document(filepath)
    #print(fp)
    string= str.replace('(CTO)', 'said')
    lower_string = string.lower()
    no_number_string = re.sub(r'\d+','',lower_string)
    no_punc_string = re.sub(r'[^\w\s]','', no_number_string) 
    #no_wspace_string = no_punc_string.strip()
    #print(no_punc_string)
    no_newline_string= re.sub('\s+',' ',no_punc_string)
    
    #no_newline_string= correct_data(no_newline_string)
    
    return no_newline_string
    #new_file_lines.append(no_wspace_string)
    #new_file_lines = []
    # for line in content:
    #     print(line)
    #     string= line.replace('(CTO)', 'said')
    #     lower_string = string.lower()
    #     no_number_string = re.sub(r'\d+','',lower_string)
    #     no_punc_string = re.sub(r'[^\w\s]','', no_number_string) 
    #     no_wspace_string = no_punc_string.strip()
    #     new_file_lines.append(no_wspace_string)
         
    # #new_file_lines= correct_data(new_file_lines)
    # return new_file_lines

def correct_data(string):
    model_name = "flexudy/t5-base-multi-sentence-doctor" # you can specify the model size here
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = T5ForConditionalGeneration.from_pretrained(model_name)
    # doc = docx.Document(filepath)
    # new_file_lines = []
    # for i in doc.paragraphs:
    print(string)
    input_text = "repair_sentence:"+string+"context: {Azur is good platfo.}{Hari is planning to present a dashboard.} </s>"

    input_ids = tokenizer.encode(input_text, return_tensors="pt")

    outputs = model.generate(input_ids, max_length=32, num_beams=1)

    sentence = tokenizer.decode(outputs[0], skip_special_tokens=True, clean_up_tokenization_spaces=True)
    #new_file_lines.append(sentence)
        
    # with open(filepath, 'w') as file1:
    #     file1.writelines(new_file_lines)
    
    return sentence
    
        #print(sentence)
#filepath= os.environ['filepath']
#normalize_text('/home/ayushri/Downloads/VirtualRoom1_2021-05-05.docx')
app.run()